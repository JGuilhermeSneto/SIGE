"""
Gera o Plano de Negócio do SIGE em formato DOCX.
Uso: python scripts/gerar_plano_negocio.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUTPUT = Path(__file__).resolve().parent.parent / "Plano_de_Negocio_SIGE.docx"


def add_title(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Capa
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PLANO DE NEGÓCIO\n\n")
    run.bold = True
    run.font.size = Pt(22)
    run2 = title.add_run("SIGE — Sistema Integrado de Gestão Escolar\n")
    run2.bold = True
    run2.font.size = Pt(16)
    run3 = title.add_run("\nEcossistema Web + Mobile + IoT para Gestão Educacional\n")
    run3.font.size = Pt(12)
    run4 = title.add_run("\nJunho de 2026\nVersão 1.0")
    run4.font.size = Pt(11)
    doc.add_page_break()

    # Introdução
    add_title(doc, "PLANO DE NEGÓCIO", 1)
    add_title(doc, "O que é e para que serve", 2)
    add_para(
        doc,
        "Este Plano de Negócio descreve a viabilidade comercial, operacional e financeira do "
        "SIGE (Sistema Integrado de Gestão Escolar), plataforma SaaS desenvolvida para "
        "escolas privadas, redes de ensino e instituições que necessitam de gestão acadêmica, "
        "administrativa, financeira e de comunicação integradas em um único ecossistema digital.",
    )
    add_para(
        doc,
        "O documento serve como mapa de percurso para captação de investidores, parceiros "
        "institucionais, programas de aceleração e planejamento estratégico da expansão "
        "comercial do produto já em estágio avançado de desenvolvimento (v8.0 Apex).",
    )

    add_title(doc, "Seu mapa de percurso", 2)
    add_para(doc, "1ª parte — A elaboração do plano de negócio: análise estratégica, mercado, "
             "marketing, operação e projeções financeiras com base no produto real.", bold=True)
    add_para(doc, "2ª parte — A construção do seu plano de negócio: modelos e roteiros "
             "consolidados para preenchimento e atualização contínua.", bold=True)

    add_title(doc, "Recomendações gerais", 2)
    add_bullets(
        doc,
        [
            "Revisar trimestralmente premissas de mercado e indicadores financeiros.",
            "Validar preços com pilotos em 3 a 5 escolas antes da escala nacional.",
            "Manter conformidade LGPD como diferencial competitivo em licitações e contratos.",
            "Priorizar módulos estáveis (acadêmico, usuários, segurança) na comercialização inicial.",
            "Documentar casos de sucesso com métricas de adoção mobile e redução de custos operacionais.",
        ],
    )
    doc.add_page_break()

    # PARTE 1
    add_title(doc, "PARTE 1 – A ELABORAÇÃO DO PLANO DE NEGÓCIO", 1)

    # 1 Sumário Executivo
    add_title(doc, "1 – Sumário Executivo", 2)
    add_title(doc, "1.1 – Resumo dos principais pontos do plano de negócio", 3)
    add_para(
        doc,
        "O SIGE é uma plataforma integrada de gestão escolar composta por quatro pilares: "
        "Backend Django (núcleo de negócio), Frontend React/Vite (gestores e professores), "
        "Aplicativo Mobile React Native/Expo (alunos e responsáveis) e camada IoT (RFID/MQTT "
        "para frequência e controle de acesso). O sistema cobre 11+ módulos funcionais, "
        "incluindo gestão acadêmica, financeira, saúde, biblioteca, comunicação, segurança "
        "(SOC/LGPD), TI/Mission Control e dashboards analíticos.",
    )
    add_para(
        doc,
        "Diferenciais técnicos comprovados: arquitetura multi-tenant, criptografia AES-256 em "
        "dados pessoais, autenticação JWT/MFA, API REST documentada (Swagger), deploy em "
        "nuvem (Render + MySQL Aiven + Cloudinary CDN), 149 casos de teste automatizados "
        "(68% de cobertura) e observabilidade com Prometheus/Grafana.",
    )
    add_para(
        doc,
        "Modelo de negócio: SaaS por assinatura mensal/anual, com planos por faixa de alunos "
        "e módulos opcionais (financeiro avançado, IoT, BI). Meta: 50 instituições clientes "
        "em 24 meses, com ticket médio de R$ 3.500/mês por escola de porte médio (200–500 alunos).",
    )

    add_title(doc, "1.2 – Dados dos empreendedores, experiência profissional e atribuições", 3)
    add_para(doc, "Empreendedor líder e responsável técnico:", bold=True)
    add_bullets(
        doc,
        [
            "J. Guilherme S. Netos — Fundador e desenvolvedor principal do ecossistema SIGE.",
            "Experiência: desenvolvimento full-stack (Python/Django, React, React Native), "
            "arquitetura de software, DevOps (Docker, Render, CI/CD) e gestão de produto digital.",
            "Atribuições: visão de produto, arquitetura técnica, desenvolvimento do core, "
            "integração mobile/API, deploy em produção e relacionamento com early adopters.",
        ],
    )
    add_para(doc, "Equipe complementar (expansão planejada):", bold=True)
    add_bullets(
        doc,
        [
            "Desenvolvedor(a) mobile — manutenção e evolução do app Expo.",
            "Analista comercial / educação — prospecção B2B em escolas privadas.",
            "Suporte e implantação — onboarding, treinamento e migração de dados.",
            "Consultor(a) jurídico-tributário — contratos, LGPD e enquadramento fiscal.",
        ],
    )

    add_title(doc, "1.3 – Dados do empreendimento", 3)
    add_bullets(
        doc,
        [
            "Nome fantasia: SIGE — Sistema Integrado de Gestão Escolar",
            "Razão social: [A definir na constituição — ex.: SIGE Tecnologia Educacional LTDA]",
            "Natureza: desenvolvimento e licenciamento de software educacional (SaaS)",
            "Sede operacional: Brasil (operação remota com infraestrutura em nuvem)",
            "Repositório e produto: github.com/JGuilhermeSneto/SIGE",
            "URL de produção: sige-g7w2.onrender.com (ambiente de demonstração/homologação)",
            "Público-alvo inicial: escolas privadas de ensino fundamental e médio (50–800 alunos)",
        ],
    )

    add_title(doc, "1.4 – Missão da empresa", 3)
    add_para(
        doc,
        "Conectar educação e tecnologia com excelência, oferecendo uma plataforma segura, "
        "acessível e integrada que simplifica a gestão escolar, aproxima famílias da vida "
        "acadêmica dos estudantes e gera dados confiáveis para decisões pedagógicas e "
        "administrativas — em total conformidade com a LGPD.",
    )

    add_title(doc, "1.5 – Setores de atividades", 3)
    add_bullets(
        doc,
        [
            "CNAE principal sugerido: 62.01-5/01 — Desenvolvimento de programas de computador sob encomenda",
            "CNAE secundário: 62.02-3/00 — Desenvolvimento e licenciamento de programas de computador customizáveis",
            "CNAE secundário: 63.19-4/00 — Portais, provedores de conteúdo e outros serviços de informação na internet",
            "Setor de mercado: EdTech / Gestão Escolar / ERP Educacional",
        ],
    )

    add_title(doc, "1.6 – Forma jurídica", 3)
    add_para(
        doc,
        "Recomendação: Sociedade Limitada (LTDA) para operação comercial B2B, contratos "
        "recorrentes e eventual entrada de sócios-investidores. Alternativa inicial: MEI apenas "
        "para validação com faturamento limitado (até teto legal), com migração para LTDA ao "
        "atingir 10+ clientes pagantes.",
    )

    add_title(doc, "1.7 – Enquadramento tributário", 3)
    add_para(doc, "Âmbito federal — Simples Nacional:", bold=True)
    add_bullets(
        doc,
        [
            "Anexo III ou V (serviços de TI), conforme enquadramento do contador.",
            "Fator R: avaliar folha vs. receita para otimização tributária.",
            "ISS municipal sobre licenciamento de software (alíquota conforme município).",
            "PIS/COFINS monofásico ou integrado ao Simples, conforme anexo.",
            "Emissão de NFS-e para cada fatura de assinatura mensal.",
        ],
    )

    add_title(doc, "1.8 – Capital social", 3)
    add_para(
        doc,
        "Capital social inicial sugerido: R$ 50.000,00 (cinquenta mil reais), integralizado em "
        "moeda corrente, distribuído entre o(s) sócio(s) fundador(es). O capital cobre "
        "formalização jurídica, infraestrutura cloud nos primeiros 12 meses, marketing "
        "inicial e reserva de caixa operacional.",
    )

    add_title(doc, "1.9 – Fonte de recursos", 3)
    add_bullets(
        doc,
        [
            "Recursos próprios do empreendedor (desenvolvimento já realizado — ativo intangível).",
            "Programas de aceleração EdTech e editais de inovação (FAPERN, SEBRAE, Finep Startup).",
            "Receita recorrente SaaS (MRR) a partir do 6º mês de comercialização ativa.",
            "Possível rodada anjo (R$ 150.000 – R$ 500.000) para equipe comercial e suporte.",
            "Parcerias com cooperativas de ensino e mantenedoras para implantação piloto.",
        ],
    )
    doc.add_page_break()

    # 2 Análise de mercado
    add_title(doc, "2 – Análise de mercado", 2)
    add_title(doc, "2.1 – Estudo dos clientes", 3)
    add_para(
        doc,
        "Segmento primário: gestores e mantenedores de escolas privadas que ainda utilizam "
        "planilhas, sistemas legados desconectados ou múltiplas ferramentas não integradas.",
    )
    add_para(doc, "Perfil do cliente ideal (ICP):", bold=True)
    add_bullets(
        doc,
        [
            "Escola privada com 100 a 600 alunos, 15 a 50 colaboradores.",
            "Dor: falta de app para pais/alunos, retrabalho administrativo, insegurança de dados.",
            "Decisor: diretor(a), mantenedor(a) ou coordenador(a) administrativo(a).",
            "Influenciadores: coordenação pedagógica, TI escolar e conselho de pais.",
            "Disposição a pagar: R$ 2.500 a R$ 6.000/mês por pacote completo web + mobile.",
        ],
    )
    add_para(doc, "Personas atendidas pelo SIGE:", bold=True)
    add_bullets(
        doc,
        [
            "Gestor escolar — cadastro, turmas, comunicados, financeiro, auditoria.",
            "Professor — notas, frequência, materiais didáticos, planejamento de aula.",
            "Aluno — boletim, frequência, mural, materiais (app mobile).",
            "Responsável — acompanhamento do filho, notificações, controle parental.",
            "Equipe de TI — Mission Control, SOC, backups, feature flags.",
        ],
    )

    add_title(doc, "2.2 – Estudo dos concorrentes", 3)
    add_para(doc, "Concorrentes diretos e indiretos no Brasil:", bold=True)
    add_bullets(
        doc,
        [
            "ERPs educacionais consolidados (TOTVS Educacional, Sophia, Sponte) — robustos, "
            "porém custos elevados e curva de implantação longa.",
            "Sistemas regionais e verticais menores — preço competitivo, mas baixa integração mobile/IoT.",
            "Planilhas + WhatsApp + Google Classroom — custo zero, sem governança nem LGPD.",
        ],
    )
    add_para(doc, "Posicionamento competitivo do SIGE:", bold=True)
    add_bullets(
        doc,
        [
            "Preço intermediário com stack moderna (API-first, mobile nativo).",
            "SOC e LGPD nativos (diferencial em contratos públicos e escolas de elite).",
            "IoT RFID para frequência automática (roadmap comercial v8+).",
            "Código aberto interno e customização por tenant (multi-escola).",
        ],
    )

    add_title(doc, "2.3 – Estudo dos fornecedores", 3)
    add_bullets(
        doc,
        [
            "Infraestrutura cloud: Render (hospedagem), Aiven (MySQL gerenciado), Cloudinary (CDN/mídia).",
            "Domínio e DNS: registradores nacionais/internacionais (Registro.br).",
            "Gateway de pagamento (futuro): Asaas, Stripe ou Pagar.me para módulo financeiro.",
            "Notificações push: Firebase Cloud Messaging / Apple Push Notification Service.",
            "Hardware IoT: fornecedores de ESP32, leitores RFID RC522 e brokers MQTT.",
            "Serviços jurídicos e contábeis: escritório parceiro para LTDA e Simples Nacional.",
        ],
    )
    doc.add_page_break()

    # 3 Marketing
    add_title(doc, "3 – Plano de Marketing", 2)
    add_title(doc, "3.1 – Descrição dos principais produtos e serviços", 3)
    add_bullets(
        doc,
        [
            "SIGE Core Web — gestão acadêmica, usuários, calendário, documentos, dashboards.",
            "SIGE Mobile — app para alunos e responsáveis (notas, frequência, perfil, push).",
            "SIGE Shield — segurança, auditoria LGPD, SOC, 2FA, monitoramento de logins.",
            "SIGE Financeiro — mensalidades, faturas, inadimplência, BI (módulo beta).",
            "SIGE Saúde — ficha médica, atestados, vacinas.",
            "SIGE Biblioteca — acervo digital e empréstimos.",
            "SIGE IoT — frequência RFID e automação (P&D, comercialização fase 2).",
            "Serviços: implantação, migração de dados, treinamento e suporte premium.",
        ],
    )

    add_title(doc, "3.2 – Preço", 3)
    add_para(doc, "Tabela de planos sugerida (valores mensais por instituição):", bold=True)
    add_bullets(
        doc,
        [
            "Plano Essencial (até 150 alunos): R$ 1.990/mês — web + mobile + acadêmico básico.",
            "Plano Profissional (151–400 alunos): R$ 3.490/mês — todos módulos estáveis + SOC.",
            "Plano Enterprise (401+ alunos): R$ 5.990/mês + R$ 8/aluno excedente — IoT + SLA prioritário.",
            "Taxa de implantação única: R$ 2.500 a R$ 8.000 conforme volume de migração.",
            "Desconto anual: 15% para pagamento antecipado (12 meses).",
        ],
    )

    add_title(doc, "3.3 – Estratégias promocionais", 3)
    add_bullets(
        doc,
        [
            "Piloto gratuito 60 dias para até 2 escolas parceiras (case de sucesso).",
            "Webinars mensais: 'Gestão escolar digital e LGPD na prática'.",
            "Presença em feiras educacionais e eventos SEBRAE/FAPERN.",
            "Marketing de conteúdo: blog técnico, comparativos e depoimentos de gestores.",
            "Programa de indicação: 1 mês grátis para cada escola indicada que fechar contrato.",
            "Demonstração ao vivo do app mobile para conselhos de pais.",
        ],
    )

    add_title(doc, "3.4 – Estrutura de comercialização", 3)
    add_bullets(
        doc,
        [
            "Canal direto B2B: prospecção ativa via LinkedIn, e-mail e visitas presenciais.",
            "Parcerias com consultorias pedagógicas e contadores de escolas.",
            "Licitações e pregões (médio prazo) com certificações LGPD documentadas.",
            "Marketplace EdTech (futuro): integração com hubs de soluções educacionais.",
            "Vendas online: landing page com agendamento de demo e onboarding self-service (fase 2).",
        ],
    )

    add_title(doc, "3.5 – Localização do negócio", 3)
    add_para(
        doc,
        "Operação digital-first, sem dependência de ponto comercial fixo na fase inicial. "
        "Infraestrutura hospedada em nuvem com CDN global (Cloudinary) e banco MySQL em "
        "Aiven (região configurável). Atendimento presencial a clientes na região de "
        "atuacao do fundador, com expansão para representantes regionais conforme carteira.",
    )
    doc.add_page_break()

    # 4 Operacional
    add_title(doc, "4 – Plano Operacional", 2)
    add_title(doc, "4.1 – Layout ou arranjo físico", 3)
    add_para(
        doc,
        "Modelo remoto/híbrido: home office para desenvolvimento e comercial; coworking "
        "para reuniões com clientes. Infraestrutura lógica (não física): monorepo com "
        "SIGE/ (Django), frontend_SIGE/ (React), SIGE_APP/ (Expo) e documentação em docs/.",
    )

    add_title(doc, "4.2 – Capacidade produtiva, comercial e de prestação de serviços", 3)
    add_bullets(
        doc,
        [
            "Capacidade técnica atual: 1 escola piloto em produção com 50+ usuários e 31 alunos ativos.",
            "Arquitetura multi-tenant: suporta múltiplas instituições no mesmo deploy.",
            "Meta operacional ano 1: até 15 escolas simultâneas sem equipe de suporte dedicada.",
            "Meta ano 2: 50 escolas com equipe de 5 pessoas (2 dev, 2 suporte, 1 comercial).",
            "SLA padrão: 99,5% uptime; SLA premium: 99,9% com suporte em até 4 horas úteis.",
        ],
    )

    add_title(doc, "4.3 – Processos operacionais", 3)
    add_bullets(
        doc,
        [
            "Onboarding: diagnóstico → migração de dados → treinamento → go-live (2–4 semanas).",
            "Desenvolvimento: Git flow, CI/CD, testes automatizados (pytest, 149 casos).",
            "Deploy: build.sh (collectstatic + migrate) no Render via GitHub.",
            "Suporte: canal por e-mail/WhatsApp + documentação em SIGE/docs/.",
            "Segurança: backups Quantum Snapshots, SOC em tempo real, revisão trimestral LGPD.",
            "Atualizações: releases mensais com changelog e comunicação aos gestores.",
        ],
    )

    add_title(doc, "4.4 – Necessidade de pessoal", 3)
    add_para(doc, "Quadro projetado — 12 meses:", bold=True)
    add_bullets(
        doc,
        [
            "1 Fundador/Dev Lead (atual) — produto e arquitetura.",
            "1 Dev Full-stack (contratação mês 6) — features e integrações.",
            "1 Analista de Sucesso do Cliente (mês 9) — onboarding e retenção.",
            "1 Vendedor B2B EdTech (mês 9) — pipeline comercial.",
            "Contador e advogado — terceirizados sob demanda.",
        ],
    )
    doc.add_page_break()

    # 5 Financeiro
    add_title(doc, "5 – Plano Financeiro", 2)
    add_title(doc, "Investimento total", 3)
    add_para(
        doc,
        "O investimento considera o ativo de software já desenvolvido (estimado em R$ 180.000 "
        "de valor de reposição em horas de engenharia) mais os aportes necessários para "
        "comercialização e operação nos primeiros 18 meses.",
    )

    add_title(doc, "5.1 – Estimativa dos investimentos fixos", 3)
    add_bullets(
        doc,
        [
            "Computadores e periféricos (2 estações): R$ 12.000",
            "Mobiliário home office: R$ 3.000",
            "Constituição LTDA + registros: R$ 4.500",
            "Identidade visual e site institucional: R$ 8.000",
            "Equipamentos IoT piloto (kits RFID): R$ 6.000",
            "Total investimentos fixos: R$ 33.500",
        ],
    )

    add_title(doc, "5.2 – Capital de giro", 3)
    add_para(doc, "Reserva de caixa para 12 meses de operação mínima:", bold=True)
    add_bullets(
        doc,
        [
            "Infraestrutura cloud (Render, Aiven, Cloudinary, domínio): R$ 800/mês × 12 = R$ 9.600",
            "Marketing e eventos: R$ 1.500/mês × 12 = R$ 18.000",
            "Pró-labore fundador (mínimo): R$ 3.000/mês × 12 = R$ 36.000",
            "Serviços terceirizados (contábil, jurídico): R$ 500/mês × 12 = R$ 6.000",
            "Imprevistos (10%): R$ 6.960",
            "Total capital de giro: R$ 76.560",
        ],
    )

    add_title(doc, "5.3 – Investimentos pré-operacionais", 3)
    add_bullets(
        doc,
        [
            "Certificações e consultoria LGPD: R$ 5.000",
            "Treinamento comercial e materiais de venda: R$ 3.000",
            "Contas de demonstração e ambientes de staging: R$ 2.000",
            "Total pré-operacional: R$ 10.000",
        ],
    )

    add_title(doc, "5.4 – Investimento total (resumo)", 3)
    add_bullets(
        doc,
        [
            "Investimentos fixos: R$ 33.500",
            "Capital de giro (12 meses): R$ 76.560",
            "Pré-operacionais: R$ 10.000",
            "INVESTIMENTO TOTAL: R$ 120.060",
            "Ativo intangível (software SIGE já desenvolvido): R$ 180.000 (valor estimado)",
        ],
    )

    add_title(doc, "5.5 – Estimativa do faturamento mensal", 3)
    add_para(doc, "Projeção conservadora de MRR (Receita Recorrente Mensal):", bold=True)
    add_bullets(
        doc,
        [
            "Meses 1–3: R$ 0 (piloto e validação)",
            "Meses 4–6: R$ 7.000 (2 escolas × R$ 3.500)",
            "Meses 7–12: R$ 24.500 (7 escolas × R$ 3.500 médio)",
            "Ano 2: R$ 87.500/mês (25 escolas × R$ 3.500 médio)",
            "Receita adicional implantação ano 1: R$ 25.000 (5 projetos × R$ 5.000)",
        ],
    )

    add_title(doc, "5.6 – Estimativa do custo unitário (infraestrutura por escola)", 3)
    add_bullets(
        doc,
        [
            "Custo marginal cloud por escola (até 400 alunos): ~R$ 45/mês",
            "Suporte e onboarding amortizado: ~R$ 120/mês por cliente no ano 1",
            "Custo variável total estimado: ~R$ 165/escola/mês",
        ],
    )

    add_title(doc, "5.7 – Estimativa dos custos de comercialização", 3)
    add_bullets(
        doc,
        [
            "Marketing digital e eventos: R$ 1.500/mês",
            "Comissão de vendas (10% MRR): variável a partir do mês 9",
            "Materiais e deslocamento: R$ 600/mês",
            "Total fixo comercial: ~R$ 2.100/mês (+ comissões)",
        ],
    )

    add_title(doc, "5.8 – Custos diretos / prestação do serviço", 3)
    add_para(
        doc,
        "Como SaaS, os custos diretos são principalmente infraestrutura cloud, APIs de "
        "terceiros e horas de implantação. Margem bruta projetada: 78% a 85% após escala "
        "de 10+ clientes.",
    )

    add_title(doc, "5.9 – Estimativa dos custos com mão de obra", 3)
    add_bullets(
        doc,
        [
            "Ano 1 (fundador + contratações parciais): R$ 72.000",
            "Ano 2 (equipe de 5): R$ 288.000",
            "Encargos e benefícios (provisão 28%): incluídos nas projeções de folha",
        ],
    )

    add_title(doc, "5.10 – Estimativa do custo com depreciação", 3)
    add_para(
        doc,
        "Depreciação anual de ativos fixos (computadores e IoT, vida útil 3 anos): "
        "R$ 33.500 / 3 ≈ R$ 11.167/ano (R$ 930/mês). Software capitalizado internamente "
        "amortizado em 5 anos: R$ 36.000/ano.",
    )

    add_title(doc, "5.11 – Estimativa dos custos fixos operacionais mensais", 3)
    add_bullets(
        doc,
        [
            "Infraestrutura cloud: R$ 800",
            "Domínio, e-mail e ferramentas (GitHub, monitoramento): R$ 350",
            "Contabilidade e jurídico: R$ 500",
            "Marketing: R$ 1.500",
            "Pró-labore / folha (média ano 1): R$ 6.000",
            "Depreciação e amortização: R$ 3.930",
            "TOTAL CUSTOS FIXOS: ~R$ 13.080/mês",
        ],
    )

    add_title(doc, "5.12 – Demonstrativo de resultados (projeção ano 1)", 3)
    add_bullets(
        doc,
        [
            "Receita bruta anual estimada: R$ 147.000",
            "(-) Custos variáveis (infra + implantação): R$ 18.000",
            "(-) Custos fixos operacionais: R$ 156.960",
            "Resultado operacional ano 1: -R$ 27.960 (esperado em fase de tração)",
            "Ponto de equilíbrio operacional: mês 14–16 (projeção)",
        ],
    )

    add_title(doc, "5.13 – Indicadores de viabilidade", 3)
    add_title(doc, "5.13.1 – Ponto de equilíbrio", 4)
    add_para(
        doc,
        "Com custos fixos de R$ 13.080/mês e ticket médio de R$ 3.500/escola (margem "
        "contribuição ~85%), o ponto de equilíbrio é atingido com aproximadamente "
        "5 escolas pagantes (R$ 17.500 MRR). Meta realista: mês 14 com pipeline comercial ativo.",
    )
    add_title(doc, "5.13.2 – Lucratividade", 4)
    add_para(
        doc,
        "Lucratividade líquida projetada no ano 2: 12% a 18% sobre receita de R$ 1.050.000, "
        "conforme retenção de clientes (churn < 8% anual) e upsell de módulos financeiro/IoT.",
    )
    add_title(doc, "5.13.3 – Rentabilidade", 4)
    add_para(
        doc,
        "ROI sobre investimento de R$ 120.060: payback estimado em 30–36 meses. "
        "Valorização do ativo de software e base recorrente de MRR elevam o valuation "
        "estimado para R$ 800.000 – R$ 1,2M ao final do ano 2 (múltiplo 8–10× MRR).",
    )
    add_title(doc, "5.13.4 – Prazo de retorno do investimento", 4)
    add_para(
        doc,
        "Prazo de retorno do investimento inicial (R$ 120.060): 30 a 36 meses. "
        "Com aporte anjo de R$ 300.000, o payback reduz para 24 meses com aceleração comercial.",
    )
    doc.add_page_break()

    # 6 Cenários
    add_title(doc, "6 – Construção de cenários", 2)
    add_para(doc, "Cenário pessimista:", bold=True)
    add_bullets(
        doc,
        [
            "3 escolas em 18 meses; MRR de R$ 10.500; necessidade de aporte adicional.",
            "Atraso no módulo financeiro reduz upsell; foco em plano Essencial.",
        ],
    )
    add_para(doc, "Cenário realista (base):", bold=True)
    add_bullets(
        doc,
        [
            "12 escolas em 18 meses; MRR de R$ 42.000; equilíbrio no mês 16.",
            "2 contratos Enterprise com IoT piloto.",
        ],
    )
    add_para(doc, "Cenário otimista:", bold=True)
    add_bullets(
        doc,
        [
            "25 escolas em 18 meses; MRR de R$ 95.000; lucro a partir do mês 12.",
            "Parceria com mantenedora regional (5 unidades simultâneas).",
            "Captação de investimento anjo para expansão nacional.",
        ],
    )
    doc.add_page_break()

    # 7 FOFA
    add_title(doc, "7 – Avaliação estratégica", 2)
    add_title(doc, "7.1 – Análise da matriz F.O.F.A (SWOT)", 3)
    add_para(doc, "Forças (Strengths):", bold=True)
    add_bullets(
        doc,
        [
            "Produto funcional em produção com web, mobile e API REST.",
            "Segurança LGPD nativa (AES-256, SOC, auditoria).",
            "Arquitetura moderna e documentação técnica extensa.",
            "Custo de infraestrutura cloud enxuto e escalável.",
        ],
    )
    add_para(doc, "Fraquezas (Weaknesses):", bold=True)
    add_bullets(
        doc,
        [
            "Equipe enxuta (dependência do fundador).",
            "Módulos financeiro e biblioteca ainda em beta.",
            "Cobertura de testes em 68% (meta 100%).",
            "Marca ainda em construção no mercado EdTech.",
        ],
    )
    add_para(doc, "Oportunidades (Opportunities):", bold=True)
    add_bullets(
        doc,
        [
            "Crescimento de escolas privadas exigindo digitalização pós-pandemia.",
            "LGPD como requisito em contratos e licitações.",
            "Demanda por apps mobile para comunicação com famílias.",
            "Integração IoT como diferencial em escolas de médio porte.",
        ],
    )
    add_para(doc, "Ameaças (Threats):", bold=True)
    add_bullets(
        doc,
        [
            "Concorrentes consolidados com força comercial.",
            "Resistência cultural à mudança em instituições tradicionais.",
            "Custos de aquisição de cliente (CAC) elevados sem parcerias.",
            "Dependência de serviços cloud terceiros (Render, Aiven).",
        ],
    )
    doc.add_page_break()

    # 8 Avaliação
    add_title(doc, "8 – Avaliação do Plano de Negócio", 2)
    add_para(
        doc,
        "O SIGE apresenta viabilidade técnica comprovada (produto em produção, banco MySQL "
        "ativo com 50+ usuários, API mobile validada) e viabilidade comercial condicionada "
        "à execução do plano de go-to-market B2B. O investimento de R$ 120.060 é compatível "
        "com o estágio MVP avançado → tração. Recomenda-se iniciar com 2 pilotos pagos, "
        "formalizar LTDA, completar módulo financeiro e buscar parceria com SEBRAE/FAPERN "
        "para reduzir risco de caixa nos primeiros 12 meses.",
    )
    add_para(doc, "Próximos passos imediatos:", bold=True)
    add_bullets(
        doc,
        [
            "Constituir LTDA e enquadrar no Simples Nacional.",
            "Fechar 2 contratos piloto com desconto de lançamento.",
            "Publicar landing page com demonstração do app mobile.",
            "Completar integração gateway de pagamento (Asaas).",
            "Elevar cobertura de testes para 85%+ nos módulos críticos.",
        ],
    )

    # 9 Roteiro
    add_title(doc, "9 – Roteiro para coleta de informações consolidado", 2)
    add_bullets(
        doc,
        [
            "Validar preços com 10 gestores escolares (entrevistas).",
            "Levantar custos reais de cloud por tenant em produção.",
            "Mapear concorrentes locais e suas tabelas de preço.",
            "Consultar contador sobre CNAE e anexo do Simples.",
            "Definir contrato padrão SaaS (SLA, LGPD, rescisão).",
            "Documentar 3 cases de uso (gestor, professor, responsável).",
            "Atualizar projeções financeiras trimestralmente.",
        ],
    )
    doc.add_page_break()

    # PARTE 2
    add_title(doc, "PARTE 2 – A CONSTRUÇÃO DO SEU PLANO DE NEGÓCIO", 1)
    add_para(
        doc,
        "Esta seção apresenta modelos resumidos para preenchimento e atualização. "
        "Os dados abaixo refletem o estado atual do projeto SIGE (junho/2026).",
    )

    sections_part2 = [
        (
            "1 – Sumário Executivo",
            [
                ("1.1 – Resumo", "SaaS EdTech integrado (web + mobile + IoT). MRR alvo R$ 42k em 18 meses."),
                ("1.2 – Empreendedores", "J. Guilherme S. Netos — fundador e tech lead."),
                ("1.3 – Empreendimento", "SIGE Tecnologia Educacional; software de gestão escolar."),
                ("1.4 – Missão", "Excelência na gestão educacional com tecnologia segura e acessível."),
                ("1.5 – Setores", "EdTech; CNAE 62.01-5/01 e correlatos."),
                ("1.6 – Forma jurídica", "LTDA (recomendado)."),
                ("1.7 – Tributário", "Simples Nacional — Anexo III/V."),
                ("1.8 – Capital social", "R$ 50.000,00"),
                ("1.9 – Fontes", "Recursos próprios, MRR, editais, possível anjo."),
            ],
        ),
        (
            "2 – Análise de mercado",
            [
                ("2.1 – Clientes", "Escolas privadas 100–600 alunos; gestores e mantenedores."),
                ("2.2 – Concorrentes", "TOTVS, Sophia, Sponte, soluções regionais."),
                ("2.3 – Fornecedores", "Render, Aiven, Cloudinary, Asaas (futuro)."),
            ],
        ),
        (
            "3 – Plano de Marketing",
            [
                ("3.1 – Produtos", "SIGE Core, Mobile, Shield, Financeiro, Saúde, Biblioteca, IoT."),
                ("3.2 – Preço", "R$ 1.990 a R$ 5.990/mês conforme porte."),
                ("3.3 – Promoção", "Piloto 60 dias, webinars, indicação, feiras."),
                ("3.4 – Comercialização", "B2B direto, parcerias, licitações (médio prazo)."),
                ("3.5 – Localização", "Digital-first; Brasil com cloud global."),
            ],
        ),
        (
            "4 – Plano Operacional",
            [
                ("4.1 – Layout", "Remoto/híbrido; monorepo SIGE + frontend + mobile."),
                ("4.2 – Capacidade", "Multi-tenant; meta 15 escolas ano 1."),
                ("4.3 – Processos", "Onboarding 2–4 sem; CI/CD; suporte documentado."),
                ("4.4 – Pessoal", "1 fundador + 3 contratações em 12 meses."),
            ],
        ),
        (
            "5 – Plano Financeiro",
            [
                ("5.1 – Fixos", "R$ 33.500"),
                ("5.2 – Giro", "R$ 76.560"),
                ("5.3 – Pré-operacional", "R$ 10.000"),
                ("5.4 – Total", "R$ 120.060"),
                ("5.5 – Faturamento", "R$ 147.000 ano 1"),
                ("5.6 – Custo unitário", "~R$ 165/escola/mês"),
                ("5.7 – Comercialização", "~R$ 2.100/mês + comissões"),
                ("5.8 – Custos diretos", "Margem bruta 78–85%"),
                ("5.9 – Mão de obra", "R$ 72.000 ano 1"),
                ("5.10 – Depreciação", "~R$ 930/mês"),
                ("5.11 – Fixos mensais", "~R$ 13.080"),
                ("5.12 – DRE", "Prejuízo operacional -R$ 27.960 ano 1 (tração)"),
                ("5.13.1 – Equilíbrio", "5 escolas pagantes"),
                ("5.13.2 – Lucratividade", "12–18% ano 2"),
                ("5.13.3 – Rentabilidade", "Valuation R$ 800k–1,2M ano 2"),
                ("5.13.4 – Payback", "30–36 meses"),
            ],
        ),
        (
            "6 – Cenários",
            [("Resumo", "Pessimista: 3 escolas | Realista: 12 | Otimista: 25 (18 meses).")],
        ),
        (
            "7 – Avaliação estratégica",
            [("7.1 – FOFA", "Força: produto pronto; Fraqueza: equipe enxuta; Oportunidade: LGPD; Ameaça: concorrência.")],
        ),
    ]

    for section_title, items in sections_part2:
        add_title(doc, section_title, 2)
        for subtitle, content in items:
            add_title(doc, subtitle, 3)
            add_para(doc, content)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "\n\nDocumento gerado automaticamente a partir da documentação técnica do SIGE.\n"
        "Fontes: README.md, COMPENDIO_TECNICO.md, ROADMAP.md, documentacao_perfis_sige.md\n"
        "Repositório: https://github.com/JGuilhermeSneto/SIGE"
    )
    run.font.size = Pt(9)
    run.italic = True

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Plano de negócio gerado: {OUTPUT}")


if __name__ == "__main__":
    main()
