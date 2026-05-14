import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_markdown_to_docx(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Handle Tables
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Skip separator line |---|---|
            if re.match(r'^\|[\-\s\|]+\|$', line):
                continue
                
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        else:
            if in_table:
                # Flush table to document
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            table.cell(i, j).text = cell_text.replace('<br>', '\n')
                in_table = False
                table_data = []

        # Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.italic = True
        elif line == '---' or line == '':
            continue
        else:
            # Regular paragraph with basic bold handling
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    convert_markdown_to_docx('Plano de teste SIGE.md', 'Plano de teste SIGE.docx')
