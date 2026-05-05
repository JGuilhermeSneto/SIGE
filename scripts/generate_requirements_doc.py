from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Documentação de Requisitos - SIGE v2.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('Este documento descreve detalhadamente os requisitos funcionais e não funcionais do Sistema Integrado de Gestão Escolar (SIGE).')

    # Functional Requirements
    doc.add_heading('1. Requisitos Funcionais (RF)', level=1)
    
    rf_data = [
        ["ID", "Módulo", "Requisito", "Referência de Código"],
        ["RF01", "Usuários", "Autenticação multi-perfil (Aluno, Professor, Gestor).", "apps.usuarios.views.autenticacao"],
        ["RF02", "Usuários", "Gestão de Perfis (dados pessoais, foto, CPF).", "apps.usuarios.views.perfis"],
        ["RF03", "Usuários", "Proteção contra Brute Force (Axes).", "config.settings.AXES_*"],
        ["RF04", "Acadêmico", "Gestão de Turmas (CRUD e enturmação).", "apps.academico.views.academico.listar_turmas"],
        ["RF05", "Acadêmico", "Vínculo de Disciplinas a Turmas e Professores.", "apps.academico.views.academico.cadastrar_disciplina_para_turma"],
        ["RF06", "Acadêmico", "Grade Horária Dinâmica por turno (Manhã/Tarde/Noite).", "apps.academico.utils.academico._get_grade_horario_turma"],
        ["RF07", "Acadêmico", "Lançamento de Notas Bimestrais (Média Automática).", "apps.academico.views.vida_escolar.lancar_nota"],
        ["RF08", "Acadêmico", "Chamada/Frequência Diária com Justificativa.", "apps.academico.views.vida_escolar.lancar_chamada"],
        ["RF09", "Acadêmico", "Atividades e Provas (Envio de arquivos).", "apps.academico.views.academico.entregar_atividade"],
        ["RF10", "Acadêmico", "Banco de Questões e Gabarito Automático.", "apps.academico.services.atividade_servico"],
        ["RF11", "Acadêmico", "Materiais Didáticos (Upload/Download).", "apps.academico.views.academico.cadastrar_editar_material"],
        ["RF12", "Acadêmico", "Histórico Escolar (Exportação PDF).", "apps.academico.views.relatorios.exportar_historico_pdf"],
        ["RF13", "Financeiro", "Emissão de Faturas de Mensalidade.", "apps.financeiro.models.Fatura"],
        ["RF14", "Financeiro", "Registro de Pagamentos e Fluxo de Caixa.", "apps.financeiro.views.listar_faturas"],
        ["RF15", "Financeiro", "Dashboard BI de Receitas vs Despesas.", "apps.financeiro.views.painel_financeiro"],
        ["RF16", "Biblioteca", "Consulta de Acervo e Disponibilidade.", "apps.biblioteca.views.acervo_biblioteca"],
        ["RF17", "Biblioteca", "Reserva e Empréstimo de Livros.", "apps.biblioteca.views.reservar_livro"],
        ["RF18", "Biblioteca", "Jornada Literária (Status de Leitura).", "apps.biblioteca.views.minhas_leituras"],
        ["RF19", "Saúde", "Ficha Médica e Registro de Vacinas.", "apps.saude.views.visualizar_saude_aluno"],
        ["RF20", "Saúde", "Gestão de Atestados com Abono de Faltas.", "apps.saude.views.revisar_atestado"],
        ["RF21", "Infraestrutura", "Gestão de Patrimônio e Inventário.", "apps.infraestrutura.views.painel_infraestrutura"],
        ["RF22", "Infraestrutura", "Controle de Estoque (Entradas/Saídas).", "apps.infraestrutura.services.InfraService"],
        ["RF23", "Comunicação", "Painel de Comunicados Institucionais.", "apps.comunicacao.views.listar_comunicados"],
        ["RF24", "Dashboards", "BI Acadêmico (Evasão e Performance).", "apps.dashboards.views"]
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(rf_data[0]):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    for row in rf_data[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    doc.add_paragraph('\n')

    # Non-Functional Requirements
    doc.add_heading('2. Requisitos Não Funcionais (RNF)', level=1)
    
    rnf_data = [
        ["ID", "Tipo", "Requisito", "Implementação Técnica"],
        ["RNF01", "Segurança", "Conformidade com LGPD (Audit Log).", "apps.seguranca.middleware.AuditMiddleware"],
        ["RNF02", "Segurança", "Autenticação MFA (2FA).", "django_otp, two_factor"],
        ["RNF03", "Segurança", "Criptografia de Dados em Repouso.", "Aiven MySQL SSL"],
        ["RNF04", "Performance", "Cache de Alta Velocidade.", "Redis (SaaS) / LocMem (Local)"],
        ["RNF05", "Performance", "Paginação Universal (Otimização DB).", "Django Paginator (select_related)"],
        ["RNF06", "Observabilidade", "Métricas de Telemetria Industrial.", "Prometheus & Grafana"],
        ["RNF07", "Observabilidade", "Monitoramento de Erros em Produção.", "Sentry / Traceback Logger"],
        ["RNF08", "Escalabilidade", "Clean Architecture (MTV + Services).", "Layered Service Architecture"],
        ["RNF09", "Resiliência", "Processamento Assíncrono (Workers).", "Celery + RabbitMQ"],
        ["RNF10", "UX/UI", "Design System Premium (Glassmorphism).", "CSS Dinâmico (Temas Azul/Cinza)"],
        ["RNF11", "Qualidade", "Cobertura de Testes Automatizados.", "Pytest (Meta 70%+)"],
        ["RNF12", "Deploy", "Infraestrutura imutável (CI/CD).", "Render Pipeline + Docker"]
    ]

    table_rnf = doc.add_table(rows=1, cols=4)
    table_rnf.style = 'Table Grid'
    hdr_cells = table_rnf.rows[0].cells
    for i, header in enumerate(rnf_data[0]):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    for row in rnf_data[1:]:
        row_cells = table_rnf.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    # Save
    doc.save('SIGE_Requisitos.docx')
    print("Arquivo SIGE_Requisitos.docx gerado com sucesso.")

if __name__ == "__main__":
    create_doc()
